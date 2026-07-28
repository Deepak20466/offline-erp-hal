"""Generic multi-format table exporters: CSV, Excel, Word, and PDF.

Every exporter takes the same shape of input — a list of column headers and a
list of row tuples/lists of already-formatted strings — so any module (clients,
contracts, invoice register, ...) can reuse them without bespoke code.
"""
import csv
import io
from typing import Sequence

import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from openpyxl.styles import Font, PatternFill
from openpyxl.utils import get_column_letter
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.units import mm
from reportlab.platypus import Image, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet
from starlette.responses import StreamingResponse

from app.utils.branding import LOGO_PATH, logo_dimensions_pt


def _content_disposition(filename: str) -> dict:
    return {"Content-Disposition": f'attachment; filename="{filename}"'}


def export_csv(filename: str, headers: Sequence[str], rows: Sequence[Sequence]) -> StreamingResponse:
    """Stream a UTF-8 CSV file."""
    buffer = io.StringIO()
    writer = csv.writer(buffer)
    writer.writerow(headers)
    writer.writerows(rows)
    buffer.seek(0)
    byte_buffer = io.BytesIO(buffer.getvalue().encode("utf-8-sig"))
    return StreamingResponse(
        byte_buffer, media_type="text/csv", headers=_content_disposition(filename)
    )


def export_excel(
    filename: str, headers: Sequence[str], rows: Sequence[Sequence], sheet_name: str = "Sheet1"
) -> StreamingResponse:
    """Stream a formatted .xlsx workbook built with pandas + openpyxl."""
    df = pd.DataFrame(rows, columns=list(headers))
    buffer = io.BytesIO()
    with pd.ExcelWriter(buffer, engine="openpyxl") as writer:
        df.to_excel(writer, index=False, sheet_name=sheet_name)
        worksheet = writer.sheets[sheet_name]
        header_fill = PatternFill(start_color="1E3A8A", end_color="1E3A8A", fill_type="solid")
        header_font = Font(color="FFFFFF", bold=True)
        for col_idx, header in enumerate(headers, start=1):
            cell = worksheet.cell(row=1, column=col_idx)
            cell.fill = header_fill
            cell.font = header_font
        for col_idx, header in enumerate(headers, start=1):
            max_len = max([len(str(header))] + [len(str(row[col_idx - 1])) for row in rows]) if rows else len(str(header))
            worksheet.column_dimensions[get_column_letter(col_idx)].width = min(max(max_len + 2, 10), 50)
    buffer.seek(0)
    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers=_content_disposition(filename),
    )


def export_word(
    filename: str, title: str, headers: Sequence[str], rows: Sequence[Sequence]
) -> StreamingResponse:
    """Stream a .docx document containing a formatted table, with the HAL letterhead logo."""
    document = Document()
    logo_paragraph = document.add_paragraph()
    logo_paragraph.add_run().add_picture(LOGO_PATH, width=Inches(1.0))
    document.add_heading(title, level=1)
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    header_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        header_cells[idx].text = str(header)
        for paragraph in header_cells[idx].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(10)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = "" if value is None else str(value)
    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers=_content_disposition(filename),
    )


def export_pdf(
    filename: str,
    title: str,
    headers: Sequence[str],
    rows: Sequence[Sequence],
    logo_path: str | None = LOGO_PATH,
    logo_width_px: int = 100,
) -> StreamingResponse:
    """Stream a landscape A4 PDF table report with the HAL letterhead logo by default."""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=landscape(A4),
        topMargin=14 * mm,
        bottomMargin=14 * mm,
        leftMargin=12 * mm,
        rightMargin=12 * mm,
    )
    styles = getSampleStyleSheet()
    elements = []

    if logo_path:
        px_to_pt = 0.75
        width, height = logo_dimensions_pt(logo_width_px * px_to_pt)
        elements.append(Image(logo_path, width=width, height=height))
        elements.append(Spacer(1, 8))

    elements.append(Paragraph(title, styles["Title"]))
    elements.append(Spacer(1, 10))

    table_data = [list(headers)] + [[("" if v is None else str(v)) for v in row] for row in rows]
    table = Table(table_data, repeatRows=1)
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1E3A8A")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 8),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#94a3b8")),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f1f5f9")]),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ]
        )
    )
    elements.append(table)
    doc.build(elements)
    buffer.seek(0)
    return StreamingResponse(buffer, media_type="application/pdf", headers=_content_disposition(filename))
